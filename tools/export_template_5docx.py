from __future__ import annotations

import re
import tempfile
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "Rendu_Simple_5PDF"
LOCAL_TEMPLATE = ROOT / "templates" / "Tempalte.docx"
DOWNLOAD_TEMPLATE = Path(r"C:\Users\Ivan\Downloads\Tempalte.docx")
TEMPLATE = LOCAL_TEMPLATE if LOCAL_TEMPLATE.exists() else DOWNLOAD_TEMPLATE

PRIMARY = "1F3864"
SECONDARY = "2E4D7B"
PALE = "F2F5FA"
TEXT = "172033"


GROUP_PARTS = [
    "17_DOSSIER_GROUPE_COMPLET_INDEX.md",
    "14_SYNTHESE_EXECUTIVE_CLIENT.md",
    "01_RAPPORT_TECHNIQUE_GROUPE.md",
    "18_SOLUTIONS_CONCRETES_DEMO.md",
    "19_ROLES_CONTRIBUTIONS_PREUVES.md",
    "Yvan FOCSA/PE-2526_M2CS_YvanFOCSA.md",
    "Youssef GUERNIOU/PE-2526_M2CS_YoussefGUERNIOU.md",
    "Kilyan FELIX/PE-2526_M2CS_KilyanFELIX.md",
    "Mahamadou DIACOUMBA/PE-2526_M2CS_MahamadouDIACOUMBA.md",
    "20_MODE_OPERATOIRE_PFSENSE_WAZUH_LAB.md",
    "21_DASHBOARDS_ALERTES_QUALIFICATION.md",
    "24_DASHBOARD_SOC_OFFLINE.md",
    "22_EXPLOITATION_VM_RUNBOOK_REX.md",
    "23_PREUVES_FINALES_CAPTURES_VIDEO_DEPOT.md",
    "25_MODE_OPERATOIRE_CAPTURE_WAZUH_PREUVES.md",
    "26_MODE_OPERATOIRE_VIDEO_DEPOT.md",
    "33_RUNBOOK_ENREGISTREMENT_VIDEO_IMMEDIAT.md",
    "27_MANIFESTE_DEPOT_ET_INTEGRITE.md",
    "28_RUNBOOK_EXPRESS_PREUVES_RESTANTES.md",
    "29_IMPORT_PREUVES_FINALES.md",
    "30_TABLEAU_BORD_STATUT_FINAL.md",
    "00_REGISTRE_EXIGENCES_ET_SYNTHESE.md",
    "02_GUIDE_DEPLOIEMENT_UTILISATION.md",
    "03_PLAYBOOKS_PROCEDURES_REX.md",
    "12_RISQUES_RGPD_CONFORMITE.md",
    "13_PLAN_RECETTE_ACCEPTATION.md",
    "07_DOSSIER_PREUVES_CAPTURES.md",
    "16_ANNEXE_CAPTURES_WAZUH.md",
    "08_AUDIT_FINAL_CONSIGNES.md",
    "05_BACKLOG_PLANNING.md",
    "10_MODE_OPERATOIRE_DEMO_JOUR_J.md",
    "11_CHECKLIST_DEPOT_FINAL.md",
]


DOCUMENTS = [
    {
        "kind": "Projet de groupe",
        "title": "Projet Daylight",
        "subtitle": "Dossier complet Cyber Trust",
        "author": "Equipe Cyber Trust",
        "role": "SOC externalise, Wazuh, pfSense, playbooks",
        "output": "00_PROJET_COMPLET_Daylight_CyberTrust.docx",
        "parts": GROUP_PARTS,
    },
    {
        "kind": "Rendu individuel",
        "title": "Yvan FOCSA",
        "subtitle": "Architecture de la solution",
        "author": "Yvan FOCSA",
        "role": "Architecte solution, pfSense, segmentation",
        "output": "01_Yvan_FOCSA_Architecte_Solution_pfSense.docx",
        "parts": ["Yvan FOCSA/PE-2526_M2CS_YvanFOCSA.md"],
    },
    {
        "kind": "Rendu individuel",
        "title": "Youssef GUERNIOU",
        "subtitle": "SIEM Wazuh et collecte",
        "author": "Youssef GUERNIOU",
        "role": "Ingenieur SIEM, Wazuh, agents, RBAC",
        "output": "02_Youssef_GUERNIOU_SIEM_Wazuh.docx",
        "parts": ["Youssef GUERNIOU/PE-2526_M2CS_YoussefGUERNIOU.md"],
    },
    {
        "kind": "Rendu individuel",
        "title": "Kilyan FELIX",
        "subtitle": "Detection, dashboards et qualification",
        "author": "Kilyan FELIX",
        "role": "Chef de projet SOC, detection, pilotage",
        "output": "03_Kilyan_FELIX_Detection_Dashboards_Qualification.docx",
        "parts": ["Kilyan FELIX/PE-2526_M2CS_KilyanFELIX.md"],
    },
    {
        "kind": "Rendu individuel",
        "title": "Mahamadou DIACOUMBA",
        "subtitle": "Exploitation lab, playbooks et REX",
        "author": "Mahamadou DIACOUMBA",
        "role": "Exploitation VM, procedures, REX incidents",
        "output": "04_Mahamadou_DIACOUMBA_Playbooks_VM_REX.docx",
        "parts": ["Mahamadou DIACOUMBA/PE-2526_M2CS_MahamadouDIACOUMBA.md"],
    },
]


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "DCE3EE", size: str = "4") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_text(paragraph, text: str, bold: bool = False, color: str | None = None, size: int | None = None) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = rgb(color)
    if size:
        run.font.size = Pt(size)


def normalize_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("__", "").replace("`", "")
    return text.strip()


def paragraph_title(source: Path) -> str:
    if not source.exists():
        return source.stem.replace("_", " ")
    text = source.read_text(encoding="utf-8", errors="ignore")
    match = re.search(r"^#\s+(.+)$", text, re.MULTILINE)
    if match:
        return normalize_inline(match.group(1))
    return source.stem.replace("_", " ")


def load_template_images() -> list[BytesIO]:
    images: list[BytesIO] = []
    if not TEMPLATE.exists():
        return images
    with ZipFile(TEMPLATE) as zf:
        for name in zf.namelist():
            if name.startswith("word/media/") and name.lower().endswith(".png"):
                images.append(BytesIO(zf.read(name)))
    return images


def configure_document(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(TEXT)

    for style_name, size, color in [
        ("Heading 1", 22, PRIMARY),
        ("Heading 2", 15, PRIMARY),
        ("Heading 3", 12, SECONDARY),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(header, "Projet Daylight · Cyber Trust", color="6B7280", size=8)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(footer, "Sup de Vinci · Mastere 2 Cybersecurite · Annee 2025-2026", color="6B7280", size=8)

    doc.core_properties.title = title
    doc.core_properties.author = "Cyber Trust"


def add_cover(doc: Document, meta: dict[str, object]) -> None:
    top = doc.add_table(rows=1, cols=2)
    top.alignment = WD_TABLE_ALIGNMENT.CENTER
    top.autofit = True
    left = top.cell(0, 0).paragraphs[0]
    add_text(left, "Mastere 2 Cybersecurite · Annee 2025-2026\n", color=SECONDARY, size=10)
    add_text(left, "Projet de groupe · Cyber Trust", bold=True, color=TEXT, size=11)
    right = top.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for image in load_template_images():
        try:
            image.seek(0)
            run = right.add_run()
            run.add_picture(image, width=Cm(1.7))
            right.add_run("  ")
        except Exception:
            continue

    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    add_text(p, str(meta["kind"]).upper(), bold=True, color=SECONDARY, size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(p, str(meta["title"]), bold=True, color=PRIMARY, size=34)

    p = doc.add_paragraph()
    add_text(p, str(meta["subtitle"]), color=TEXT, size=18)

    doc.add_paragraph()
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cells = [
        ("Etablissement", "Sup de Vinci · Mastere 2 Cybersecurite"),
        ("Client", "Daylight"),
        ("Auteur", str(meta["author"])),
        ("Role", str(meta["role"])),
        ("Equipe", "Y. FOCSA · Y. GUERNIOU · K. FELIX · M. DIACOUMBA"),
        ("Version", "1.1 · Mise en page finale"),
    ]
    for idx, (label, value) in enumerate(cells):
        cell = table.cell(idx // 2, idx % 2)
        set_cell_shading(cell, PALE)
        set_cell_border(cell, "FFFFFF", "0")
        paragraph = cell.paragraphs[0]
        add_text(paragraph, f"{label}\n", color="6B7280", size=8)
        add_text(paragraph, value, bold=True, color=TEXT, size=10)

    for _ in range(6):
        doc.add_paragraph()
    bottom = doc.add_paragraph()
    bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(bottom, "Prestataire fictif : Cyber Trust · Date de rendu : 26/06/2026", color="4B5563", size=9)
    doc.add_page_break()


def add_summary(doc: Document, meta: dict[str, object]) -> None:
    doc.add_heading("Sommaire", level=1)
    p = doc.add_paragraph()
    add_text(p, "Rendu structure pour la correction, la soutenance et la consultation rapide des preuves.")
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).text = "Type de rendu"
    table.cell(0, 1).text = str(meta["kind"])
    row = table.add_row()
    row.cells[0].text = "Perimetre"
    row.cells[1].text = str(meta["role"])
    set_table_borders(table)
    for row in table.rows:
        set_cell_shading(row.cells[0], PALE)

    doc.add_paragraph()
    parts = list(meta["parts"])  # type: ignore[index]
    for idx, part in enumerate(parts, start=1):
        source = ROOT / part
        p = doc.add_paragraph(style="List Number")
        add_text(p, paragraph_title(source))
    doc.add_page_break()


def split_table_row(line: str) -> list[str]:
    cells = [normalize_inline(cell) for cell in line.strip().strip("|").split("|")]
    return cells


def is_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*", line))


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx in range(width):
            cell = table.cell(r_idx, c_idx)
            text = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            paragraph = cell.paragraphs[0]
            add_text(paragraph, text, bold=(r_idx == 0), color=PRIMARY if r_idx == 0 else TEXT, size=8 if width > 4 else 9)
            if r_idx == 0:
                set_cell_shading(cell, PALE)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "FAFBFD")
    doc.add_paragraph()


def resolve_image(source: Path, target: str) -> Path | None:
    clean = target.strip().strip("<>").split()[0]
    candidates = [(source.parent / clean).resolve(), (ROOT / clean).resolve()]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def add_image(doc: Document, source: Path, line: str) -> bool:
    match = re.match(r"!\[[^\]]*\]\(([^)]+)\)", line.strip())
    if not match:
        return False
    path = resolve_image(source, match.group(1))
    if not path:
        return False
    try:
        doc.add_picture(str(path), width=Inches(6.2))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    except Exception:
        return False


def add_paragraph_from_markdown(doc: Document, text: str) -> None:
    text = normalize_inline(text)
    if not text:
        return
    paragraph = doc.add_paragraph()
    add_text(paragraph, text)


def add_source(doc: Document, part: str, single: bool) -> None:
    source = ROOT / part
    raw = source.read_text(encoding="utf-8", errors="ignore")
    raw = re.sub(r"```mermaid\n(.*?)```", r"```\n\1\n```", raw, flags=re.DOTALL)
    lines = raw.splitlines()

    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []
    first_h1_skipped = False

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            add_markdown_table(doc, table_rows)
            table_rows = []

    def flush_code() -> None:
        nonlocal code_lines
        if code_lines:
            p = doc.add_paragraph()
            set_paragraph_shading(p, "F4F6FA")
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(8.5)
            code_lines = []

    if not single:
        doc.add_heading(paragraph_title(source), level=1)

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_table()
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not stripped:
            flush_table()
            continue

        if stripped.startswith("|"):
            if is_separator(stripped):
                continue
            table_rows.append(split_table_row(stripped))
            continue

        flush_table()

        if add_image(doc, source, stripped):
            continue

        if stripped.startswith("# "):
            title = normalize_inline(stripped[2:])
            if single and not first_h1_skipped:
                first_h1_skipped = True
                continue
            doc.add_heading(title, level=1)
        elif stripped.startswith("## "):
            doc.add_heading(normalize_inline(stripped[3:]), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(normalize_inline(stripped[4:]), level=3)
        elif re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_text(p, normalize_inline(re.sub(r"^[-*]\s+", "", stripped)))
        elif re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_text(p, normalize_inline(re.sub(r"^\d+\.\s+", "", stripped)))
        elif stripped.startswith(">"):
            p = doc.add_paragraph()
            set_paragraph_shading(p, PALE)
            add_text(p, normalize_inline(stripped.lstrip("> ")), color=TEXT)
        else:
            add_paragraph_from_markdown(doc, stripped)

    flush_table()
    if in_code:
        flush_code()


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def export_document(meta: dict[str, object]) -> Path:
    doc = Document()
    configure_document(doc, str(meta["title"]))
    add_cover(doc, meta)
    add_summary(doc, meta)

    parts = list(meta["parts"])  # type: ignore[index]
    single = len(parts) == 1
    for idx, part in enumerate(parts):
        if idx:
            doc.add_section(WD_SECTION_START.NEW_PAGE)
        add_source(doc, part, single=single)

    output = OUT_DIR / str(meta["output"])
    doc.save(output)
    return output


def main() -> int:
    OUT_DIR.mkdir(exist_ok=True)
    with tempfile.TemporaryDirectory():
        for meta in DOCUMENTS:
            path = export_document(meta)
            print(f"[OK] {path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
